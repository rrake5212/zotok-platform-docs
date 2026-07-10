#!/usr/bin/env python3
"""Generate ZoTok User Manual — Home Dashboard & Threads (DOCX)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = "/mnt/d/AgentWork/zotok user manual"
FRAME_DIR = OUT_DIR

doc = Document()

# ── Page Setup ──
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_before = Pt(4)
style.paragraph_format.space_after = Pt(4)

for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F) if level == 1 else (RGBColor(0x2B, 0x57, 0x9A) if level == 2 else RGBColor(0x33, 0x33, 0x33))
    hs.font.bold = True
    sizes = {1: 18, 2: 15, 3: 13}
    hs.font.size = Pt(sizes[level])
    sp = {1: (18, 12), 2: (14, 10), 3: (10, 8)}
    hs.paragraph_format.space_before = Pt(sp[level][0])
    hs.paragraph_format.space_after = Pt(sp[level][1])

# ── Helpers ──

def add_para(text, bold=False, italic=False, size=10.5, color="333333", align=None, spacing_after=6, spacing_before=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = { "center": WD_ALIGN_PARAGRAPH.CENTER, "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
                        "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT }.get(align)
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
    run.bold = bold
    run.italic = italic
    return p

def add_spacer(pts=8):
    return add_para("", size=8, spacing_after=pts, spacing_before=0)

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def add_img(file_name, width=5.5):
    path = os.path.join(FRAME_DIR, file_name)
    if not os.path.exists(path):
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    return p

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color="333333", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])

def add_header_row(table, cells):
    row = table.rows[0]
    for i, (text, width) in enumerate(cells):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_text(cell, text, bold=True, color="FFFFFF", size=9)
        set_cell_shading(cell, "1E3A5F")

def add_data_row(table, cells, shade=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_text(cell, text, bold=(i == 0), color="333333" if not shade else "444444", size=9)
        if shade:
            set_cell_shading(cell, "F5F7FA")

def add_tip_box(text):
    add_spacer(4)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Use a table with left border for tip box effect
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "EBF0F7")
    # Set left border
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="2B579A"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="2B579A"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="2B579A"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="2B579A"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(tc_borders)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"💡 Tip: ")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    add_spacer(4)

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════
add_spacer(40)
add_para("ZoTok", bold=True, size=26, color="1E3A5F", align="center", spacing_after=4)
add_para("User Manual", size=20, color="2B579A", align="center", spacing_after=2)
add_para("Home Dashboard & Threads", size=14, color="555555", align="center", spacing_after=6)
add_para("─" * 50, size=10, color="AAAAAA", align="center", spacing_after=8)
add_para("Version 1.0  •  July 2026", size=10, color="888888", align="center", spacing_after=2)
add_para("A Complete Walkthrough for New Users", size=10, color="888888", align="center")
add_spacer(30)
add_para("Hd Agencies", bold=True, size=14, color="1E3A5F", align="center", spacing_before=20)
add_para("Powered by ZoTok", size=9, color="999999", align="center")
add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Introduction",
    "   1.1 About This Manual",
    "   1.2 Logging In",
    "2. Home Dashboard",
    "   2.1 Navigation Sidebar",
    "   2.2 KPI Cards",
    "   2.3 Customer Data Table",
    "3. Settings & Configurations",
    "   3.1 Sales View Settings",
    "   3.2 Payment View Settings",
    "4. Product Module",
    "   4.1 Product KPIs",
    "   4.2 Empty State & Adding Products",
    "5. Threads",
    "   5.1 Threads Interface",
    "   5.2 Thread Types",
    "   5.3 Empty Conversation State",
    "6. Queries Management",
    "   6.1 Queries List",
    "   6.2 Query Type Details",
    "   6.3 Adding a New Query Type",
    "   6.4 Complete Query Type Reference",
    "7. Appendix",
    "   7.1 Click Flow Diagram",
    "   7.2 Tips & Shortcuts",
    "   7.3 Document Information",
]
for item in toc_items:
    indent = 0
    t = item.strip()
    if item.startswith("   "):
        indent = 1
        t = item.strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.3 * indent)
    r = p.add_run(t)
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if indent == 0:
        r.bold = True

add_page_break()

# ══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
add_para(
    'ZoTok is a comprehensive field operations and distributor management platform '
    'designed to streamline communication, sales tracking, inventory management, and '
    'payment reconciliation for businesses. This user manual covers two core modules: '
    'the Home Dashboard and the Threads (conversation/query management) feature.',
    align="justified"
)

doc.add_heading('1.1 About This Manual', level=2)
add_para(
    'This manual is structured to follow the natural workflow of a ZoTok user. '
    'Each chapter corresponds to a screen or feature demonstrated in the walkthrough, '
    'providing descriptions of UI elements, configuration options, and actionable steps.',
    align="justified"
)

doc.add_heading('1.2 Logging In', level=2)
add_para(
    'After successful authentication, users are greeted with the Home Dashboard displaying '
    'the firm name (e.g., "Hello, Hd Agencies") and a welcome message. '
    'The navigation sidebar provides access to all core modules.',
    align="justified"
)
add_img("01_Dashboard_Home_00000.jpg", width=4.5)
add_tip_box("The Home Dashboard is your command center — all modules are accessible from the left sidebar.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 2. HOME DASHBOARD
# ══════════════════════════════════════════════════════════
doc.add_heading('2. Home Dashboard', level=1)
add_para(
    'The Home Dashboard serves as the landing page after login. It provides a high-level '
    'overview of key business metrics and quick access to all ZoTok modules.',
    align="justified"
)

doc.add_heading('2.1 Navigation Sidebar', level=2)
add_para('The left sidebar contains the following navigation items:')

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([1.0, 1.0, 4.5]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Module", 1.0), ("Icon", 1.0), ("Description", 4.5)])
nav_items = [
    ("Home", "🏠", "Dashboard overview with KPI cards, receivables ageing, top customer data"),
    ("Sales", "💰", "Sales transactions, charts, column configuration"),
    ("Payments", "💳", "Payment tracking, age-wise receivables, ageing configuration"),
    ("Item", "📦", "Product catalog, inventory tracking, stock management"),
    ("Dashboard", "📊", "Advanced analytics and charts"),
    ("Day Book", "📖", "Daily transaction register"),
    ("Threads", "💬", "Customer conversations, queries, and signal inbox"),
]
for i, (mod, icon, desc) in enumerate(nav_items):
    add_data_row(t, [(mod, 1.0), (icon, 1.0), (desc, 4.5)], shade=(i % 2 == 1))

add_spacer(6)
add_para('Click any navigation item to switch to the corresponding module.')

doc.add_heading('2.2 KPI Cards', level=2)
add_para('The top of the dashboard displays key performance indicators for "This Month Till Date":')
add_bullet("Receivable — Total outstanding receivables")
add_bullet("Ageing Buckets — 0-30 Days, 31-60 Days, Above 60 Days")
add_bullet("Total Sales — Aggregated month-to-date sales figure")

doc.add_heading('2.3 Customer Data Table', level=2)
add_para('The main content area displays a sortable data table with columns including:')
add_bullet("Customer / Firm Name and Customer Name")
add_bullet("District, State, Customer Code, Phone, GSTIN")
add_bullet("Total Sales, Category Name/Code, Average Sales")
add_bullet("Total Invoices, Total Cgst")
add_tip_box("Click column headers to sort. Use the settings overlay to toggle visible columns.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 3. SETTINGS & CONFIGURATIONS
# ══════════════════════════════════════════════════════════
doc.add_heading('3. Settings & Configurations', level=1)
add_para(
    'ZoTok provides per-module configuration panels that can be opened by clicking on '
    'navigation items. These overlays allow users to customize their view.',
    align="justified"
)

doc.add_heading('3.1 Sales View Settings', level=2)
add_para('Triggered by clicking the Sales nav item, this overlay provides:')
add_bullet("Chart Configuration — Chart 1 & 2 selection, Donut Chart toggle")
add_bullet("Table Columns — Check/uncheck columns (customerFirmName, customerName, District, State, customerCode, phone, gstin, Total Sales, Category, Average Sales, Total Invoices, Total Cgst)")
add_bullet("Data Filters — Type, Customer, Category")
add_para('A state-wise summary table shows aggregated sales (e.g., Telangana: ₹900+).')
add_img("02_Sidebar_Menu_Open_00405.jpg", width=4.5)

doc.add_heading('3.2 Payment View Settings', level=2)
add_para('Triggered by clicking the Payments nav item, this overlay configures receivable ageing:')
add_bullet("Ageing By — Choose Due Date or Bill Date")
add_para("Ageing Bucket Ranges:")

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([1.5, 2.0, 3.0]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Bucket", 1.5), ("Range", 2.0), ("Description", 3.0)])
buckets = [
    ("1", "0 - 10 Days", "Current / near-term receivables"),
    ("2", "11 - 20 Days", "Early overdue"),
    ("3", "21 - 30 Days", "Moderate overdue"),
    ("4", "31 - 40 Days", "Attention required"),
    ("5", "41 - 50 Days", "Escalation needed"),
    ("6", "Over 50 Days", "Critical — immediate follow-up"),
]
for i, (b, r, d) in enumerate(buckets):
    add_data_row(t, [(b, 1.5), (r, 2.0), (d, 3.0)], shade=(i % 2 == 1))

add_spacer(4)
add_para('When no invoices exist, the system displays: "No invoices to show".', italic=True)

add_page_break()

# ══════════════════════════════════════════════════════════
# 4. PRODUCT MODULE
# ══════════════════════════════════════════════════════════
doc.add_heading('4. Product Module', level=1)
add_para(
    'The Product (Item) module allows businesses to manage their product catalog and '
    'track inventory, revenue, and sales performance.',
    align="justified"
)

doc.add_heading('4.1 Product KPIs', level=2)
add_para('Upon entering the Product module, these KPI cards are displayed:')
add_bullet("Current Stock — Total units currently in stock")
add_bullet("Inventory Value — Monetary value of current inventory (₹)")
add_bullet("Total Revenue — Revenue generated from product sales")
add_bullet("Units Sold — Total units sold")

doc.add_heading('4.2 Empty State & Adding Products', level=2)
add_para('If no products have been added yet, the system shows:')
add_para('"No products added yet. Add products to start tracking sales insights."', italic=True, align="center")
add_img("06_Navigation_to_Module_02175.jpg", width=4.5)
add_para('Click the "Add Product" button to populate your catalog.')
add_tip_box("Add all products before tracking sales to ensure accurate KPI calculations.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 5. THREADS
# ══════════════════════════════════════════════════════════
doc.add_heading('5. Threads', level=1)
add_para(
    'The Threads module is ZoTok\'s conversation management system. It organizes all '
    'customer interactions — queries, follow-ups, issues, and signals — in a unified inbox.',
    align="justified"
)

doc.add_heading('5.1 Threads Interface', level=2)
add_para('Upon navigating to Threads, the interface displays:')
add_bullet("Header showing 'Threads' with customer count (e.g., 8 Customers)")
add_img("09_Form_Entry_03105.jpg", width=4.5)

add_para("Tabbed navigation:", bold=True)
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([1.5, 5.0]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Tab", 1.5), ("Purpose", 5.0)])
tabs = [
    ("Groups", "Group conversations (broadcast to multiple customers)"),
    ("Threads", "Individual conversation threads with customers"),
    ("Issues", "Escalated tickets and problem reports"),
    ("Signal Inbox", "Incoming automated signals and alerts"),
]
for i, (tab, purpose) in enumerate(tabs):
    add_data_row(t, [(tab, 1.5), (purpose, 5.0)], shade=(i % 2 == 1))

add_spacer(4)
add_para("Category filters:", bold=True)
add_bullet("All Categories — Show all conversation types")
add_bullet("Unactioned — Show only conversations requiring action")

doc.add_heading('5.2 Thread Types', level=2)
add_para('Common conversation categories include:')
add_bullet("New Query Flow — First-time customer inquiries")
add_bullet("Bulk Order Negotiation — Large order discussions")
add_bullet("Payment Follow-up — Outstanding payment reminders")
add_bullet("Product Catalogue / Browsing Request — Catalog sharing")
add_bullet("Company Info / About Request — Business information queries")
add_bullet("Conversation Initiation — Simple greeting / onboarding")

doc.add_heading('5.3 Empty Conversation State', level=2)
add_para('When no conversation is selected, the system displays:')
add_para('"Select a conversation from the list to view messages"', italic=True, align="center")
add_tip_box("Conversations are automatically categorized. Ensure query types are configured for accurate routing.")

add_page_break()

# ══════════════════════════════════════════════════════════
# 6. QUERIES MANAGEMENT
# ══════════════════════════════════════════════════════════
doc.add_heading('6. Queries Management', level=1)
add_para(
    'The Queries Management section provides a comprehensive list of all configured query types. '
    'Each query type defines a conversation category and its associated action items.',
    align="justified"
)

doc.add_heading('6.1 Queries List', level=2)
add_para('The queries table displays configured queries (45 total, paginated 20 per page):')

t = doc.add_table(rows=1, cols=6)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [0.5, 2.0, 2.5, 1.0, 1.0, 0.8]
for i, w in enumerate(widths):
    t.columns[i].width = Inches(w)
add_header_row(t, [
    ("S.No", 0.5), ("Query Name & Description", 2.0), ("Action Items", 2.5),
    ("Last Used", 1.0), ("Updated On", 1.0), ("Actions", 0.8)
])
sample_queries = [
    ("1", "Bulk Order Negotiation", "Create Order", "19 min ago", "6 Jul 2026", "✏️"),
    ("2", "Conversation Initiation — Simple Greeting", "Request Payment", "19 min ago", "6 Jul 2026", "✏️"),
    ("3", "Create Order", "Upload Invoice, Create Order, Send Order Receipt", "19 min ago", "6 Jul 2026", "✏️"),
    ("4", "Payment Follow-up", "Send Payment Receipt", "22 hrs ago", "5 Jul 2026", "✏️"),
    ("5", "Product Catalogue / Browsing Request", "N.A.", "1 hr ago", "6 Jul 2026", "✏️"),
]
for i, row_data in enumerate(sample_queries):
    data = list(zip(row_data, widths))
    add_data_row(t, data, shade=(i % 2 == 1))

add_spacer(4)
add_img("11_Summary_Screen_06780.jpg", width=4.5)

doc.add_heading('6.2 Query Type Details Overlay', level=2)
add_para('Clicking a query row opens details showing:')
add_bullet("Query Name — The category name")
add_bullet("Description — What this query type covers")
add_bullet("Action Items — Available actions (Create Order, Send Invoice, Request Payment, etc.)")

doc.add_heading('6.3 Adding a New Query Type', level=2)
add_para('Click the "Add Query" button to open the creation form.')
add_img("12_Action_Confirmation_06930.jpg", width=4.5)

add_para("Form Fields:", bold=True)
t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([1.5, 1.0, 4.0]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Field", 1.5), ("Required", 1.0), ("Description", 4.0)])
add_data_row(t, [("Query Name *", 1.5), ("Yes", 1.0), ("A unique name for the query type", 4.0)])
add_data_row(t, [("Prompt *", 1.5), ("Yes", 1.0), ("Instruction/prompt associated with this query", 4.0)], shade=True)
add_data_row(t, [("Actions", 1.5), ("Optional", 1.0), ("Select from available message actions", 4.0)])

add_spacer(4)
add_para('Buttons: XCancel (discard) | Save (add query)', bold=True, align="center")

doc.add_heading('6.4 Complete Query Type Reference', level=2)
add_para('Pre-configured query types in the system:')

t = doc.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
widths2 = [0.5, 2.5, 2.5, 1.0]
for i, w in enumerate(widths2):
    t.columns[i].width = Inches(w)
add_header_row(t, [("#", 0.5), ("Query Name", 2.5), ("Action Items", 2.5), ("Status", 1.0)])
all_queries = [
    ("1", "Bulk Order Negotiation", "Create Order", "Active"),
    ("2", "Conversation Initiation — Simple Greeting", "Request Payment", "Active"),
    ("3", "Create Order", "Upload Invoice, Create Order, Send Order Receipt", "Active"),
    ("4", "Test create of signal", "N.A.", "Active"),
    ("5", "Product Catalogue / Browsing Request", "N.A.", "Active"),
    ("6", "Company Info / About Request", "N.A.", "Active"),
    ("7", "New Query Flow", "N.A.", "Active"),
    ("8", "Payment Follow-up", "Send Payment Receipt", "Active"),
    ("9", "Location Sharing / GPS Coordinates", "Send Ledger, Send Invoices", "Active"),
    ("10", "Location Sharing for Pickup/Delivery", "Send Template", "Active"),
    ("11", "Catalogue-led Order Placement", "Send Payment Receipt", "Active"),
    ("12", "Field Location Pings / Geo Check-ins", "N.A.", "Active"),
    ("13", "No stock — not available with us", "N.A.", "Active"),
    ("14", "Link/URL Share", "N.A.", "Active"),
    ("15", "Short Bot Commands / Quick Requests", "Payment Reminder, Send Invoice", "Active"),
    ("16", "Change Address", "N.A.", "Active"),
    ("17", "Share Catalog", "Create Order", "Active"),
]
for i, row_data in enumerate(all_queries):
    data = list(zip(row_data, widths2))
    add_data_row(t, data, shade=(i % 2 == 1))

add_page_break()

# ══════════════════════════════════════════════════════════
# 7. APPENDIX
# ══════════════════════════════════════════════════════════
doc.add_heading('7. Appendix', level=1)

doc.add_heading('7.1 Click Flow Diagram', level=2)
add_para('Complete navigation flow from the walkthrough video:')

t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([0.8, 5.7]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Step", 0.8), ("Action", 5.7)])
steps = [
    ("1", "Login → Home Dashboard (Hello, Hd Agencies)"),
    ("2", "Click Sales → Sales View Settings overlay (charts, columns)"),
    ("3", "Save/Cancel → Return to Home Dashboard"),
    ("4", "Click Payments → Payment View Settings (ageing config)"),
    ("5", "Save/Cancel → Return to Home Dashboard"),
    ("6", "Click Item → Product View Settings → Products Home"),
    ("7", "Click Add Product → Product creation flow"),
    ("8", "Navigate to Threads → Threads interface (8 Customers)"),
    ("9", "Select thread → Conversation view or empty state"),
    ("10", "Open Queries → Queries list (45 types, 2 pages)"),
    ("11", "Click Add Query → Query Type form (Name + Prompt)"),
    ("12", "Save → Query added to list"),
]
for i, row_data in enumerate(steps):
    add_data_row(t, list(zip(row_data, [0.8, 5.7])), shade=(i % 2 == 1))

add_spacer(8)
add_para("Flow Diagram:", bold=True)
add_para("""
Home Dashboard
  ├── Sales View Settings → Save/Cancel → Home
  ├── Payments → Payment View Settings → Save/Cancel → Home
  ├── Item → Product View Settings → Products Home → Add Product
  └── Threads → Threads Interface (8 Customers, tabs)
               ├── Select Thread → Queries List (45 queries)
               └── Add Query → Query Type Form → Save
""", size=9, color="555555")

doc.add_heading('7.2 Tips & Shortcuts', level=2)
add_bullet("Use sidebar navigation to switch between modules instantly")
add_bullet("Settings overlays can be dismissed by clicking Cancel or pressing Escape")
add_bullet("Tables support sorting by clicking column headers")
add_bullet("Use category filters in Threads to narrow down conversation types")
add_bullet("Pagination controls at bottom of lists allow page navigation")
add_bullet("Configure query types upfront for automatic conversation routing")
add_bullet("Regularly review Unactioned threads to avoid missing customer queries")

doc.add_heading('7.3 Document Information', level=2)
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, w in enumerate([2.5, 4.0]):
    t.columns[i].width = Inches(w)
add_header_row(t, [("Field", 2.5), ("Value", 4.0)])
info = [
    ("Product", "ZoTok (Field Operations & Distributor Management)"),
    ("Module", "Home Dashboard & Threads"),
    ("Document Version", "1.0 (July 2026)"),
    ("Source Video", "Home and Threads.mp4 (4 min, 213 MB)"),
    ("Video Resolution", "1916 × 900 px, 30 fps, H.264"),
    ("Firm (Demo)", "Hd Agencies"),
    ("Frame Extraction Tool", "OpenCV + RapidOCR"),
    ("Document Format", "Microsoft Word (.docx)"),
]
for i, (field, value) in enumerate(info):
    add_data_row(t, [(field, 2.5), (value, 4.0)], shade=(i % 2 == 1))

# ── Save ──
out_path = os.path.join(OUT_DIR, "ZoTok_User_Manual_Home_and_Threads.docx")
doc.save(out_path)
file_size = os.path.getsize(out_path)
print(f"✅ DOCX created: {out_path}")
print(f"   Size: {file_size / 1024:.1f} KB")
